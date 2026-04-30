import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

# Đường dẫn gốc là thư mục chứa file này (DACN1/)
BASE_PATH = r"D:\HocTap\HK6\DACN1"

DATASET_DIR = os.path.join(BASE_PATH, "dataset")
LEGAL_FORMS_DIR = os.path.join(DATASET_DIR, "legal_forms")
GENERATED_FORMS_DIR = os.path.join(DATASET_DIR, "generated_forms")


class LegalFormGenerator:
    def __init__(self):
        # Tạo thư mục nếu chưa có
        os.makedirs(LEGAL_FORMS_DIR, exist_ok=True)
        os.makedirs(GENERATED_FORMS_DIR, exist_ok=True)

    def _create_surrender_docx(self, species_vn, species_latin, legal_group, folder_name, user_data=None):
        """Tạo file Word .docx theo mẫu bạn cung cấp"""
        if user_data is None:
            # Giá trị mặc định (bạn có thể thay sau khi thêm form nhập thông tin cá nhân)
            user_data = {
                'city': 'Nha Trang',
                'fullName': '.......................................................................................................',
                'idCard': '...........................',
                'idCardDate': '................',
                'idCardPlace': '................',
                'address': '...........................................................................................',
                'phone': '.......................................................................................',
                'department': 'Chi cục Kiểm lâm tỉnh Khánh Hòa',
                'healthStatus': '...........................................................................',
            }

        doc = Document()

        # --- Cấu hình font chữ chuẩn hành chính ---
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(13)

        # --- 1. Quốc hiệu ---
        header_qh = doc.add_paragraph()
        header_qh.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_qh = header_qh.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        run_qh.bold = True

        header_tn = doc.add_paragraph()
        header_tn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_tn = header_tn.add_run("Độc lập - Tự do - Hạnh phúc")
        run_tn.bold = True
        doc.add_paragraph("-" * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- 2. Ngày tháng ---
        now = datetime.datetime.now()
        date_line = doc.add_paragraph()
        date_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_line.add_run(f"{user_data['city']}, ngày {now.day} tháng {now.month} năm {now.year}")

        # --- 3. Tiêu đề đơn ---
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = title.add_run("\nĐƠN TỰ NGUYỆN BÀN GIAO ĐỘNG VẬT HOANG DÃ")
        run_title.bold = True
        run_title.font.size = Pt(16)

        # --- 4. Kính gửi ---
        kinh_gui = doc.add_paragraph()
        kinh_gui.add_run("Kính gửi: ").bold = True
        kinh_gui.add_run(f"Chi cục Kiểm lâm tỉnh/thành phố {user_data['department']}")

        # --- 5. Thông tin người làm đơn ---
        doc.add_paragraph(f"Tôi tên là: {user_data['fullName']}")
        doc.add_paragraph(f"Số CCCD/CMND: {user_data['idCard']}  Ngày cấp: {user_data['idCardDate']}  Nơi cấp: {user_data['idCardPlace']}")
        doc.add_paragraph(f"Địa chỉ thường trú: {user_data['address']}")
        doc.add_paragraph(f"Số điện thoại liên lạc: {user_data['phone']}")

        # --- 6. Nội dung bàn giao ---
        doc.add_paragraph("\nNay tôi làm đơn này tự nguyện bàn giao cho quý cơ quan cá thể động vật sau:")

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Tên loài'
        hdr_cells[1].text = 'Tên khoa học'
        hdr_cells[2].text = 'Phân loại pháp lý'

        row_cells = table.add_row().cells
        row_cells[0].text = species_vn
        row_cells[1].text = species_latin
        row_cells[2].text = legal_group

        # --- 7. Cam kết ---
        doc.add_paragraph(f"\nTình trạng sức khỏe hiện tại: {user_data['healthStatus']}")
        doc.add_paragraph("Lý do bàn giao: Tôi nhận thức được đây là loài động vật nguy cấp quý hiếm cần được bảo tồn và tự nguyện bàn giao cho cơ quan chức năng để cứu hộ, tái thả về tự nhiên.")
        doc.add_paragraph("Tôi cam đoan thông tin trên là đúng sự thật và hoàn toàn chịu trách nhiệm trước pháp luật.")

        # --- 8. Chữ ký ---
        sign_section = doc.add_table(rows=1, cols=2)
        sign_section.columns[0].width = doc.sections[0].page_width // 4
        sign_section.columns[1].width = doc.sections[0].page_width // 4

        sign_cell = sign_section.rows[0].cells[1]
        p_title = sign_cell.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.add_run("NGƯỜI LÀM ĐƠN\n").bold = True
        p_title.add_run("(Ký, ghi rõ họ tên)\n\n\n\n\n")

        p_name = sign_cell.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.add_run(user_data['fullName']).bold = True

        # Lưu file
        file_name = f"Don_Ban_Giao_{folder_name}.docx"
        output_path = os.path.join(GENERATED_FORMS_DIR, file_name)
        doc.save(output_path)
        print(f"--- [THÀNH CÔNG] Đã tạo file Word: {output_path}")
        return file_name

    def generate_form(self, template_name, data, output_filename=None):
        """
        data chứa: ten_loai_tieng_anh, ten_viet_nam, nhom_phap_ly
        """
        species_latin = data.get("ten_loai_tieng_anh", "Không rõ")
        species_vn = data.get("ten_viet_nam", "Không rõ")
        legal_group = data.get("nhom_phap_ly", "Không rõ")
        folder_name = species_latin.replace(" ", "_").replace("/", "_")

        return self._create_surrender_docx(species_vn, species_latin, legal_group, folder_name)

    def preview_form(self, template_name, data):
        """Trả về nội dung text để hiển thị trong modal Xem trước"""
        species_latin = data.get("ten_loai_tieng_anh", "Không rõ")
        species_vn = data.get("ten_viet_nam", "Không rõ")
        legal_group = data.get("nhom_phap_ly", "Không rõ")

        preview_text = f"""ĐƠN TỰ NGUYỆN BÀN GIAO ĐỘNG VẬT HOANG DÃ

Kính gửi: Chi cục Kiểm lâm tỉnh/thành phố ...

Tôi tên là: .................................................................
Số CCCD/CMND: ................   Ngày cấp: ........   Nơi cấp: ........
Địa chỉ thường trú: .................................................................
Số điện thoại: .................................................................

Nay tôi làm đơn này tự nguyện bàn giao cá thể động vật sau:

Tên loài          : {species_vn}
Tên khoa học      : {species_latin}
Phân loại pháp lý : {legal_group}

Tình trạng sức khỏe: ........................................................
Lý do bàn giao: Tôi nhận thức được đây là loài động vật nguy cấp quý hiếm...

Tôi cam đoan thông tin trên là đúng sự thật và hoàn toàn chịu trách nhiệm trước pháp luật.

NGƯỜI LÀM ĐƠN
(Ký, ghi rõ họ tên)

.................................................................
"""
        return preview_text